from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Program\AI Program")
PROJECT = ROOT / "do-agent"
TEMPLATE = ROOT / "模板.docx"
OUT = PROJECT / "软件工程综合课程实习报告-Dodo-Agent通用智能体平台.docx"


def set_east_asia_font(run, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    set_east_asia_font(r)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="8EAADB"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if r_idx == 0:
                shade_cell(cell, "D9EAF7")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        set_east_asia_font(run, "黑体")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_para(doc, text="", style=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if first_line and text:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        set_east_asia_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• " + item)
        r.font.size = Pt(10.5)
        set_east_asia_font(r)


def add_numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{idx}. {item}")
        r.font.size = Pt(10.5)
        set_east_asia_font(r)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)
    set_east_asia_font(r)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
    style_table(table, widths)
    doc.add_paragraph()
    return table


def replace_cover_text(doc):
    mapping = {
        "组长：                学号：            手机号：": "组长：________________    学号：________________    手机号：________________",
        "组员： \t  学号：": "组员：________________    学号：________________",
        "学号：": "组员：________________    学号：________________",
        "成绩Grade：": "成绩 Grade：",
    }
    seen_member_line = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "学号：":
            seen_member_line += 1
            p.text = "组员：________________    学号：________________"
        elif text in mapping:
            p.text = mapping[text]
        for r in p.runs:
            set_east_asia_font(r)


def clear_after_catalog(doc):
    found = False
    body = doc._body._body
    for p in list(doc.paragraphs):
        if found:
            body.remove(p._element)
        elif p.text.strip() == "目录":
            found = True


def add_static_catalog(doc):
    entries = [
        ("1 项目概述", 1), ("1.1 编写目的", 2), ("1.2 项目背景", 2), ("1.3 问题定义", 2), ("1.4 可行性研究", 2), ("1.5 开发环境", 2),
        ("2 需求分析", 1), ("2.1 需求调研", 2), ("2.2 业务需求", 2), ("2.3 系统功能性需求", 2), ("2.4 非功能性需求", 2),
        ("3 系统设计", 1), ("3.1 概要设计", 2), ("3.2 数据库设计", 2), ("3.3 接口设计", 2), ("3.4 用户界面设计", 2), ("3.5 详细设计", 2),
        ("4 系统实现", 1), ("4.1 后端实现", 2), ("4.2 文件与向量检索实现", 2), ("4.3 PPT 生成实现", 2), ("4.4 前端实现", 2),
        ("5 软件测试", 1), ("5.1 测试概述", 2), ("5.2 功能测试", 2), ("5.3 接口测试", 2), ("5.4 部署测试", 2),
        ("6 软件发布与法律问题", 1), ("7 小组分工及总结", 1),
    ]
    for title, level in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7 if level == 2 else 0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.font.size = Pt(10.5)
        if level == 1:
            r.bold = True
        set_east_asia_font(r)
    doc.add_page_break()


def add_image_if_exists(doc, path, width_cm, caption):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)


def build():
    doc = Document(str(TEMPLATE))
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    for style_name in ["Normal"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
    existing_styles = {s.name for s in doc.styles}
    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        if style_name in existing_styles:
            style = doc.styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(31, 78, 121)

    replace_cover_text(doc)
    clear_after_catalog(doc)
    add_static_catalog(doc)

    add_heading(doc, "1 项目概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本报告依据软件工程综合课程实习的要求，对 Dodo-Agent 企业级端到端通用智能体平台进行需求分析、概要设计、详细设计、实现说明、测试设计和发布合规说明。报告重点说明项目的工程结构、核心算法流程、数据库设计、接口设计以及可运行部署方式，为课程验收、后续维护和二次开发提供完整文档依据。")
    add_heading(doc, "1.2 项目背景", 2)
    add_para(doc, "大语言模型应用正在从单轮问答走向能够理解任务、调用工具、检索资料并持续执行的智能体系统。传统聊天系统通常只完成文本生成，难以覆盖企业知识库问答、网页实时检索、文件分析、自动生成演示文稿、复杂研究任务拆解等场景。Dodo-Agent 以 Spring Boot 和 Spring AI 为工程底座，将 ReAct、Plan-Execute、RAG、MCP 工具调用和模板驱动 PPT 生成集成到统一平台中，形成面向实际业务的通用智能体项目。")
    add_heading(doc, "1.3 问题定义", 2)
    add_para(doc, "本项目要解决的核心问题是：如何在一个可部署、可扩展、可观测的 Web 系统中，将大模型推理能力与外部工具、企业文件、向量数据库和业务状态持久化结合起来，使用户能够通过自然语言完成信息查询、文件问答、深度研究和办公文档生成。")
    add_bullets(doc, [
        "统一入口问题：不同类型智能体需要通过一致的流式协议向前端输出思考过程、文本答案、工具状态、引用来源和推荐问题。",
        "上下文与记忆问题：系统需要保存会话历史，并在多轮任务中控制上下文长度，避免模型调用失效或成本过高。",
        "工具执行问题：智能体需要安全地调用搜索、文件读取、文件系统、Grep、Bash、Skills 等工具，并处理并发、失败和用户中断。",
        "知识检索问题：上传文件既要保存原始对象，也要进行解析、切分、向量化和基于文件 ID 的隔离检索。",
        "业务闭环问题：PPT 生成等长流程任务需要状态机、断点续传、结果持久化和文件托管能力。"
    ])
    add_heading(doc, "1.4 可行性研究", 2)
    add_para(doc, "技术方面，项目采用 JDK 21、Spring Boot 3.5.6、Spring AI 1.1.0、MyBatis-Plus、Project Reactor、MySQL、Redis、PostgreSQL + pgvector、MinIO 等成熟组件，核心依赖可通过 Maven 管理，具有较高可行性。经济方面，项目主要依赖开源框架和可配置的大模型接口，部署成本集中在模型调用和中间件资源上，适合课程项目和企业原型验证。操作方面，系统提供 Web 前端，用户通过选择 Agent、上传文件、输入自然语言即可使用，学习成本较低。法律方面，项目通过环境变量管理密钥，避免将敏感配置提交仓库；生成内容、联网搜索结果和文件处理需要在实际发布时补充用户隐私协议、第三方模型服务协议和内容审核策略。")
    add_heading(doc, "1.5 开发环境", 2)
    add_table(doc, ["类别", "环境与工具", "说明"], [
        ["后端语言", "Java 21", "主工程语言，使用 Spring Boot 启动和管理 Bean"],
        ["构建工具", "Maven 3.9+", "依赖管理、编译、打包和运行"],
        ["Web 框架", "Spring Boot 3.5.6", "REST 接口、静态资源托管和配置管理"],
        ["AI 框架", "Spring AI 1.1.0 / Spring AI Alibaba", "模型调用、工具调用、RAG 和向量存储集成"],
        ["数据库", "MySQL 8.0", "会话、文件元数据、PPT 实例和模板持久化"],
        ["向量库", "PostgreSQL + pgvector", "大文件语义切片与相似度检索"],
        ["缓存/任务", "Redis / Redisson", "为并发控制和分布式扩展预留能力"],
        ["对象存储", "MinIO", "上传文件与生成 PPT 文件托管"],
        ["前端", "Vue + 原生 JS/CSS", "聊天界面、文件上传、SSE 流式渲染"],
    ], [3.0, 4.2, 7.0])

    add_heading(doc, "2 需求分析", 1)
    add_heading(doc, "2.1 需求调研", 2)
    add_para(doc, "调研对象包括课程实践中的智能问答系统、企业知识库问答系统、自动化办公工具和深度研究类智能体。现有系统常见痛点包括：功能分散、无法中断长任务、缺少工具执行过程展示、上传文件与会话关联弱、生成类任务缺少状态管理等。Dodo-Agent 将这些需求归纳为“多智能体协同 + 统一流式交互 + 持久化业务状态”的系统目标。")
    add_heading(doc, "2.2 业务需求", 2)
    add_table(doc, ["参与者", "业务目标", "典型操作"], [
        ["普通用户", "通过自然语言获取答案和执行任务", "发起网页搜索、选择深度研究、上传文件问答、生成 PPT"],
        ["知识工作者", "提高资料整理和办公产出效率", "基于报告生成摘要、检索内部文件、生成演示文稿"],
        ["系统维护者", "保证系统稳定运行和可扩展", "配置模型密钥、中间件连接、Skills 目录和搜索服务"],
        ["开发者", "扩展 Agent 与工具能力", "继承 BaseAgent、注册 ToolCallback、增加新的状态策略"],
    ], [3.0, 5.2, 6.0])
    add_heading(doc, "2.3 系统功能性需求", 2)
    add_table(doc, ["模块", "功能需求", "优先级"], [
        ["智能问答", "支持 WebSearchReactAgent 使用 Tavily MCP 进行联网搜索，并返回引用来源", "高"],
        ["文件问答", "支持 PDF、Word、TXT、PNG、JPG 上传、解析、存储和基于 fileId 的问答", "高"],
        ["深度研究", "支持需求澄清、研究主题生成、计划拆解、多轮执行、自我批判和最终报告生成", "高"],
        ["PPT 生成", "支持 CREATE、MODIFY、RESUME 三类意图，基于模板和状态机生成 PPT", "高"],
        ["Skills 智能体", "整合搜索、文件、Skills、文件系统、Grep、Bash 等工具，由模型按需选择", "高"],
        ["会话管理", "支持会话列表、详情查询、删除及跨轮历史记忆恢复", "高"],
        ["任务控制", "支持基于 conversationId 的运行任务去重和停止", "中"],
        ["前端交互", "支持 Agent 切换、文件拖拽上传、SSE 流式显示、思考过程与引用折叠", "中"],
    ], [3.0, 8.0, 2.2])
    add_heading(doc, "2.4 非功能性需求", 2)
    add_bullets(doc, [
        "性能需求：流式接口应尽快返回 thinking 或 text 事件；文件向量化采用批量写入，减少向量库请求次数。",
        "可靠性需求：AgentTaskManager 统一记录运行任务和 Disposable，支持用户中断后释放资源。",
        "可维护性需求：不同智能体继承 BaseAgent，公共能力包括会话记忆、响应封装、计时、工具记录和推荐问题生成。",
        "扩展性需求：通过 Spring AI ToolCallback、MCP Client 和 Skills 目录扩展工具，不需要改动前端协议。",
        "安全性需求：密钥通过环境变量注入，上传文件以 UUID 作为 fileId，数据库记录和对象存储路径分离。",
        "兼容性需求：前端使用标准 EventSource/fetch stream 思路处理 SSE 数据，后端采用 UTF-8 输出。"
    ])

    add_heading(doc, "3 系统设计", 1)
    add_heading(doc, "3.1 概要设计", 2)
    add_para(doc, "系统采用前后端一体化部署方式，后端由 Spring Boot 提供 REST/SSE 接口和静态资源托管，前端负责会话、文件上传和流式消息渲染。智能体层以 BaseAgent 为抽象父类，向下调用模型、工具和业务服务；服务层处理文件、会话、PPT 实例、对象存储和向量检索；数据层由 MySQL、pgvector 和 MinIO 组成。")
    add_image_if_exists(doc, PROJECT / "docs" / "architecture" / "dodo-agent-architecture.png", 14.0, "图 3-1 Dodo-Agent 整体架构图")
    add_table(doc, ["层次", "主要组件", "职责"], [
        ["表现层", "index.html、app.js、api.js、style.css", "对话界面、文件上传、SSE 解析、Markdown 渲染、会话列表"],
        ["控制层", "AgentController、FileController、SessionController", "暴露智能体、文件和会话接口，完成参数校验和服务调度"],
        ["智能体层", "WebSearchReactAgent、FileReactAgent、PlanExecuteAgent、PPTBuilderAgent、SkillsReactAgent", "组织提示词、工具调用、流式输出、状态推进和结果汇总"],
        ["服务层", "FileManageService、EmbeddingService、AiSessionService、PptPythonRenderService、MinioService", "业务处理、向量化、持久化、对象存储和渲染"],
        ["基础设施层", "MySQL、PostgreSQL + pgvector、Redis、MinIO、MCP、LLM API", "数据、向量、缓存、文件、工具和模型能力支撑"],
    ], [2.8, 5.0, 7.2])
    add_heading(doc, "3.1.1 功能模块设计", 3)
    add_table(doc, ["模块", "输入", "处理", "输出"], [
        ["Web 搜索问答", "query、conversationId", "加载历史记忆，调用 Tavily MCP，ReAct 多轮推理", "text、thinking、reference、recommend"],
        ["文件问答", "query、conversationId、fileId", "读取 fileId，调用 loadContent 工具，必要时走 RAG 检索", "基于文件内容的答案"],
        ["深度研究", "复杂问题", "澄清需求、生成主题、计划执行、批判、上下文压缩", "结构化研究报告与引用"],
        ["PPT 生成", "PPT 需求或修改需求", "意图识别、模板选择、搜索、生成大纲与 Schema、渲染、上传", "PPT 下载链接和过程说明"],
        ["Skills 通用智能体", "自然语言任务、可选 fileId", "模型自主选择 Skills、搜索、文件系统、Grep、Bash 等工具", "通用任务执行结果"],
    ], [3.0, 3.8, 5.8, 4.0])
    add_heading(doc, "3.1.2 核心流程设计", 3)
    add_para(doc, "深度研究流程体现了本项目的智能体自主执行能力。PlanExecuteAgent 先判断用户问题是否需要补充信息，再将问题转化为研究主题；随后按轮次生成计划，并按任务 order 分组并发执行。每个任务内部使用 SimpleReactAgent 调用工具完成信息检索，轮次结束后由 critique 方法评估信息是否充分，未通过则继续迭代，通过后进入 summarizeStream 生成最终报告。")
    add_image_if_exists(doc, PROJECT / "docs" / "architecture" / "deep-research-engine.png", 14.0, "图 3-2 深度研究引擎流程图")
    add_numbered(doc, [
        "需求澄清：通过 REQUIREMENT_CLARIFICATION 提示词判断信息是否充分，不充分时暂停并要求用户补充。",
        "主题生成：将原始问题转化为更明确的研究主题，作为后续计划生成的上下文。",
        "计划拆解：模型输出 PlanTask 列表，任务包含 id、instruction 和 order，order 相同的任务可并行执行。",
        "工具执行：使用 Semaphore 将工具并发控制在 3 个以内，并记录搜索引用和工具结果。",
        "结果批判：根据当前计划与工具结果判断研究是否通过，不通过则进入下一轮。",
        "上下文压缩：当状态字符数超过阈值时，生成压缩快照以控制上下文规模。",
        "报告汇总：提取工具结果，生成最终文本，并发送 reference 事件给前端。"
    ])
    add_heading(doc, "3.2 数据库设计", 2)
    add_para(doc, "系统使用 MySQL 保存业务状态和会话数据，使用 pgvector 保存文件切片向量。MySQL 表结构集中在 ai_session、ai_file_info、ai_ppt_inst 和 ai_ppt_template 四张核心表，覆盖用户交互历史、文件元数据、PPT 生成实例和 PPT 模板定义。")
    add_table(doc, ["数据表", "主键/关键字段", "用途"], [
        ["ai_session", "id、session_id、question、answer、thinking、tools、reference、agent_type、fileid", "保存用户问题、AI 答案、思考过程、工具记录、引用和推荐问题"],
        ["ai_file_info", "id、file_id、file_name、file_type、minio_path、extracted_text、status、embed", "保存上传文件的元数据、解析文本、处理状态和是否向量化"],
        ["ai_ppt_inst", "id、conversation_id、template_code、status、requirement、search_info、outline、ppt_schema、file_url", "保存 PPT 生成流程中的状态和中间产物，支持断点续传"],
        ["ai_ppt_template", "id、template_code、template_schema、file_path、style_tags、slide_count", "保存可选 PPT 模板及其页面字段约束"],
        ["vector_file_info", "document id、embedding、metadata(fileid, chunkId)", "pgvector 表，按 fileid 过滤检索文件切片"],
    ], [3.0, 6.2, 5.0])
    add_heading(doc, "3.3 接口设计", 2)
    add_table(doc, ["接口", "方法", "参数", "说明"], [
        ["/agent/chat/stream", "GET SSE", "query、conversationId", "联网搜索智能问答"],
        ["/agent/file/stream", "GET SSE", "query、conversationId、fileId", "基于上传文件的问答"],
        ["/agent/pptx/stream", "GET SSE", "query、conversationId", "PPT 生成、修改或继续"],
        ["/agent/deep/stream", "GET SSE", "query、conversationId", "深度研究任务"],
        ["/agent/skills/stream", "GET SSE", "query、conversationId、fileId?", "通用 Skills 智能体"],
        ["/agent/stop", "GET", "conversationId", "停止指定会话正在执行的任务"],
        ["/file/upload", "POST", "file", "上传文件并返回 fileId"],
        ["/file/info/{fileId}", "GET", "fileId", "查询文件元数据"],
        ["/session/{conversationId}", "GET", "conversationId", "查询会话详情"],
        ["/session/list", "GET", "pageNum、pageSize", "分页查询会话列表"],
    ], [3.4, 2.4, 4.0, 5.0])
    add_heading(doc, "3.4 用户界面设计", 2)
    add_para(doc, "前端采用 Vue 组合式 API 管理状态。页面核心区域包括左侧会话列表、Agent 选择区、消息展示区、文件上传区和输入框。用户选择不同 Agent 后，前端根据 selectedAgent 拼接不同流式接口；上传文件成功后保存 uploadedFileId，在文件问答或 Skills 模式下随请求一起发送。")
    add_bullets(doc, [
        "消息展示：用户消息和助手消息分角色展示，助手消息支持 Markdown 与代码高亮。",
        "思考过程：thinking、tool_start、tool_end、error 被加入 timeline，可折叠查看。",
        "引用来源：reference 事件被解析为引用列表，支持后续验证答案来源。",
        "推荐问题：recommend 事件生成快捷追问，用户点击后自动填入并发送。",
        "中断执行：发送过程中可调用 /agent/stop 停止后端任务。"
    ])
    add_heading(doc, "3.5 详细设计", 2)
    add_heading(doc, "3.5.1 BaseAgent 抽象设计", 3)
    add_para(doc, "BaseAgent 封装了所有智能体共享的能力，包括持久化 ChatMemory 创建、历史加载、统一 AgentResponse JSON 构造、首次响应时间和总耗时统计、运行任务检查、任务注册、工具记录以及推荐问题生成。子类只需要实现 execute 方法即可接入统一协议。")
    add_heading(doc, "3.5.2 文件 RAG 设计", 3)
    add_para(doc, "FileManageService 在文件上传时先生成 UUID fileId 并写入 ai_file_info，随后上传 MinIO。PDF、DOC、DOCX、TXT 会通过 FileParserService 解析文本；文本长度超过 5000 字符时，系统使用 OverlapParagraphTextSplitter 按 500 字符、50 字符重叠切分为 Document，并附加 fileid 和 chunkId 元数据后批量写入 pgvector。EmbeddingService 检索时先压缩用户问题，再扩展为 3 个查询，并以 fileid 过滤相似度搜索结果，保证不同文件之间的数据隔离。")
    add_heading(doc, "3.5.3 PPT 状态机设计", 3)
    add_image_if_exists(doc, PROJECT / "docs" / "architecture" / "ppt-builder-engine.png", 14.0, "图 3-3 PPT 智能生成引擎")
    add_para(doc, "PPTBuilderAgent 通过 PptIntentRecognizer 识别 CREATE_PPT、MODIFY_PPT 和 RESUME_PPT 三类意图。创建流程会生成 AiPptInst，并交由 PptStateStrategyFactory 根据实例状态依次执行模板选择、需求澄清、联网搜索、大纲生成、Schema 生成、渲染和成功处理。所有中间状态写入 ai_ppt_inst，因此用户中断或刷新后可通过 RESUME_PPT 从当前状态继续。")
    add_heading(doc, "3.5.4 Skills 通用智能体设计", 3)
    add_para(doc, "SkillsReactAgent 将 SkillsTool、FileContentService、FileSystemTools、GrepTool、BashTool 和联网搜索工具合并为 ToolCallback 数组，模型在 ReAct 轮次中自主选择工具。该 Agent 使用 AgentStreamEvent 输出 Thinking、Text、ToolStart、ToolEnd、Error 和 Complete，前端可展示完整工具轨迹。ContextCompactor 在每轮调用前检查上下文策略，必要时压缩历史消息，提升长任务稳定性。")

    add_heading(doc, "4 系统实现", 1)
    add_heading(doc, "4.1 后端实现", 2)
    add_para(doc, "后端入口为 DodoAgentApplication，使用 @SpringBootApplication 启动。application.yml 中配置端口 8888、DashScope 兼容 OpenAI 接口、MySQL、Redis、MinIO、pgvector 和 Tavily MCP。AgentController 在 afterPropertiesSet 中初始化 Tavily MCP ToolCallback，并按请求动态构造对应 Agent。")
    add_table(doc, ["类名", "实现要点"], [
        ["AgentController", "提供五类 Agent 的 SSE 接口，完成工具合并、持久化记忆加载和任务停止"],
        ["BaseAgent", "封装会话记忆、响应类型、计时、任务检查、工具记录和推荐问题"],
        ["AgentTaskManager", "按 conversationId 管理运行任务、Sink 和 Disposable，支持去重与停止"],
        ["AgentResponse / AgentStreamEvent", "定义统一流式响应协议，便于前端按类型渲染"],
        ["GlobalExceptionHandler", "统一异常处理，提升接口错误输出一致性"],
    ], [4.0, 10.5])
    add_heading(doc, "4.2 文件与向量检索实现", 2)
    add_para(doc, "文件上传接口 /file/upload 接收 MultipartFile，校验非空后调用 FileManageService.uploadFile。服务层完成文件元数据保存、MinIO 上传、文本解析或图片识别、大文件向量化等步骤。图片文件使用 qwen3-vl-plus 多模态模型生成纯文本描述，使图片也可以纳入后续问答。")
    add_table(doc, ["步骤", "实现类/方法", "说明"], [
        ["1", "FileController.uploadFile", "接收文件并返回 FileInfo"],
        ["2", "FileManageService.uploadFile", "生成 fileId，保存 PROCESSING 状态，上传 MinIO"],
        ["3", "FileParserService.parseFile", "解析 PDF、Word、TXT 文本"],
        ["4", "processLargeFileEmbedding", "大文件切分并添加 fileid 元数据"],
        ["5", "EmbeddingService.embedAndStore", "按 9 条一批写入 pgvector"],
        ["6", "EmbeddingService.ragRetrieve", "问题压缩、查询扩展、fileid 过滤相似度检索"],
    ], [1.5, 4.8, 8.2])
    add_heading(doc, "4.3 PPT 生成实现", 2)
    add_para(doc, "PPT 生成模块由 Java 状态机和 Python 渲染服务共同完成。数据库中的 ai_ppt_template 保存模板字段 Schema，PPTBuilderAgent 生成符合模板约束的 PptSchema，RenderStrategy 调用 PptPythonRenderService 将 Schema 渲染为 PPT 文件，再通过 MinIO 返回可访问地址。修改 PPT 时，系统读取已有 ppt_schema，结合用户修改需求重新生成 Schema 并进入渲染流程。")
    add_heading(doc, "4.4 前端实现", 2)
    add_para(doc, "前端 app.js 使用 Vue ref 和 computed 管理状态，sendMessage 方法根据 selectedAgent 构造流式请求 URL，并使用 fetch + ReadableStream 逐行解析 SSE 数据。processStreamData 按类型更新消息内容、思考时间线、工具状态、引用来源和推荐问题。文件上传通过 APP_API.uploadFile 完成，上传成功后保存 fileId 并展示文件名。")
    add_heading(doc, "4.5 关键工程特性实现", 2)
    add_bullets(doc, [
        "统一流式协议：后端所有 Agent 输出 JSON 字符串，字段 type 表示 thinking、text、reference、recommend、error 或 complete。",
        "持久化记忆：每次请求先将历史记录转换为 ChatMemory，保证跨接口的多轮上下文连续。",
        "任务中断：用户停止时，AgentTaskManager 停止对应 Disposable 并从任务表移除。",
        "引用溯源：深度研究收集 SimpleReactAgent 的 SearchResult，在最终答案后发送 reference。",
        "断点续传：PPT 任务状态持久化在 ai_ppt_inst，按 status 重新进入状态机。"
    ])

    add_heading(doc, "5 软件测试", 1)
    add_heading(doc, "5.1 测试概述", 2)
    add_para(doc, "测试目标是验证系统的主要功能链路、接口参数校验、长任务中断、文件处理和状态持久化。由于本项目依赖外部模型、Tavily MCP、MinIO、MySQL、pgvector 等服务，课程测试可分为本地单服务测试、集成环境测试和人工交互测试三个层次。")
    add_heading(doc, "5.2 功能测试", 2)
    add_table(doc, ["编号", "测试项", "输入/操作", "预期结果"], [
        ["TC01", "普通智能问答", "选择 chat，输入实时问题", "页面流式显示思考与答案，数据库保存会话"],
        ["TC02", "深度研究", "选择 deep，输入复杂调研任务", "经历需求分析、计划执行、批判和总结阶段"],
        ["TC03", "文件上传", "上传 PDF/DOCX/TXT", "返回 fileId，ai_file_info 状态为 SUCCESS"],
        ["TC04", "大文件问答", "上传超过 5000 字文本并提问", "文件被切片向量化，检索结果与 fileId 匹配"],
        ["TC05", "图片问答", "上传 PNG/JPG 并提问", "多模态模型生成图片描述，回答围绕图片内容"],
        ["TC06", "PPT 创建", "输入“生成一份 AI 科技主题 PPT”", "生成 PPT 实例，最终返回文件地址"],
        ["TC07", "PPT 修改", "在同会话输入修改要求", "读取已有 Schema，重新生成并渲染 PPT"],
        ["TC08", "任务停止", "长任务执行时点击停止", "后端停止对应 conversationId 的任务，前端不再继续流式追加"],
        ["TC09", "会话列表", "刷新页面并查询历史", "可加载历史会话和消息详情"],
    ], [1.5, 3.0, 5.5, 5.0])
    add_heading(doc, "5.3 接口测试", 2)
    add_table(doc, ["接口", "异常输入", "预期处理"], [
        ["/agent/chat/stream", "query 为空", "返回参数不能为空错误"],
        ["/agent/file/stream", "fileId 为空", "返回文件 ID 不能为空错误"],
        ["/file/upload", "空文件", "返回“文件不能为空”"],
        ["/file/info/{fileId}", "不存在的 fileId", "返回文件不存在或错误结果"],
        ["/session/{conversationId}", "不存在的会话", "返回会话不存在"],
        ["/agent/stop", "无运行任务的 conversationId", "返回未找到正在执行任务"],
    ], [4.0, 5.0, 5.0])
    add_heading(doc, "5.4 部署测试", 2)
    add_numbered(doc, [
        "准备 JDK 21、Maven 3.9+、MySQL 8.0、Redis、PostgreSQL + pgvector 和 MinIO。",
        "根据 .env.example 配置 DASHSCOPE_API_KEY、OPENAI_BASE_URL、TAVILY_API_KEY、MYSQL_URL、PGVECTOR_*、MINIO_* 等变量。",
        "执行 src/main/resources/sql/ai_db.sql 初始化 MySQL 表结构。",
        "运行 mvn clean package 检查编译，运行 mvn spring-boot:run 启动服务。",
        "访问 http://localhost:8888，完成问答、文件上传和会话查询的冒烟测试。"
    ])

    add_heading(doc, "6 软件发布与法律问题", 1)
    add_heading(doc, "6.1 发布方式", 2)
    add_para(doc, "项目可作为独立 Spring Boot 应用发布。发布前应将 .env.example 作为配置说明，实际密钥通过服务器环境变量或密钥管理系统注入；MySQL、Redis、pgvector 和 MinIO 建议使用 Docker Compose 或云托管服务统一部署。生产环境需要启用 HTTPS、访问鉴权、上传文件大小限制、日志脱敏和服务健康检查。")
    add_heading(doc, "6.2 法律与合规问题", 2)
    add_bullets(doc, [
        "隐私保护：上传文件可能包含个人或企业敏感信息，应明确告知用户文件用途、保存周期和删除方式。",
        "第三方服务：模型接口、Tavily 搜索、MinIO 或云数据库需要遵循相应服务条款。",
        "内容安全：生成内容可能存在不准确或不合规风险，发布时应增加免责声明和人工审核机制。",
        "版权风险：用户上传资料和生成 PPT 所使用的图片、文本、模板应确保授权范围明确。",
        "安全责任：Bash、文件系统等高权限工具在生产环境必须做权限隔离、目录限制和审计记录。"
    ])

    add_heading(doc, "7 小组分工及总结", 1)
    add_heading(doc, "7.1 小组成员分工", 2)
    add_table(doc, ["成员", "主要分工", "成果"], [
        ["成员 A", "总体架构与后端基础框架", "Spring Boot 项目结构、配置、统一响应和异常处理"],
        ["成员 B", "智能体核心逻辑", "BaseAgent、WebSearchReactAgent、PlanExecuteAgent、SkillsReactAgent"],
        ["成员 C", "文件与知识库模块", "文件上传、MinIO 存储、文本解析、向量化和 RAG 检索"],
        ["成员 D", "PPT 生成模块", "PPTBuilderAgent、状态策略、模板 Schema、Python 渲染"],
        ["成员 E", "前端与测试文档", "Vue 聊天界面、SSE 渲染、测试用例和课程报告整理"],
    ], [3.0, 5.0, 6.0])
    add_heading(doc, "7.2 团队总结", 2)
    add_para(doc, "通过本次课程实践，团队完成了一个具有工程完整性的通用智能体平台原型。项目不仅包含传统 Web 系统的控制器、服务、数据库和前端页面，还引入了大模型应用特有的工具调用、流式输出、RAG、上下文压缩和状态机任务编排。开发过程中最大的收获是认识到智能体系统并非简单调用模型接口，而是需要围绕任务状态、工具边界、数据持久化、异常恢复和用户体验进行系统化设计。")
    add_heading(doc, "7.3 个人心得体会", 2)
    add_para(doc, "本项目加深了我对软件工程“需求-设计-实现-测试-发布”完整流程的理解。尤其是在阅读和整理 Dodo-Agent 代码时，可以清楚看到抽象设计的重要性：BaseAgent 降低了多类智能体的重复成本，状态机提升了 PPT 长流程任务的可恢复性，统一流式协议让前后端协作更加稳定。后续如果继续完善该项目，可以重点补充权限认证、自动化测试、容器化部署、工具沙箱和更细粒度的审计能力，使其从课程项目进一步走向可生产化的智能体平台。")
    add_heading(doc, "参考文献", 1)
    refs = [
        "Spring Boot 官方文档：https://spring.io/projects/spring-boot",
        "Spring AI 官方文档：https://docs.spring.io/spring-ai/reference/",
        "MyBatis-Plus 官方文档：https://baomidou.com/",
        "MinIO 官方文档：https://min.io/docs/",
        "PostgreSQL pgvector 项目文档：https://github.com/pgvector/pgvector",
        "Model Context Protocol 官方资料：https://modelcontextprotocol.io/",
        "Dodo-Agent 项目源码与 README 文档。"
    ]
    add_numbered(doc, refs)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Dodo-Agent 通用智能体平台软件工程综合课程实习报告"
        for r in footer.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(127, 127, 127)
            set_east_asia_font(r)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
