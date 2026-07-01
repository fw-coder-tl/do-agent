from pathlib import Path
import json
import re
from collections import Counter

import pdfplumber
from docx import Document


ROOT = Path(r"D:\Program\AI Program")
PROJECT = ROOT / "do-agent"
PDF = ROOT / "13组 赵晨璐2206010206 软件综合实践报告.pdf"
TEMPLATE = ROOT / "模板.docx"
OUT = PROJECT / "tmp_report_analysis"
OUT.mkdir(exist_ok=True)


def docx_summary(path: Path):
    doc = Document(str(path))
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "style": p.style.name if p.style else "",
                "text": text[:300],
            })
    tables = []
    for ti, t in enumerate(doc.tables):
        rows = []
        for r in t.rows[:8]:
            rows.append([c.text.strip()[:120] for c in r.cells])
        tables.append({"index": ti, "rows": rows})
    sections = []
    for s in doc.sections:
        sections.append({
            "page_width": s.page_width.twips,
            "page_height": s.page_height.twips,
            "top_margin": s.top_margin.twips,
            "bottom_margin": s.bottom_margin.twips,
            "left_margin": s.left_margin.twips,
            "right_margin": s.right_margin.twips,
        })
    return {"paragraphs": paragraphs, "tables": tables, "sections": sections}


def pdf_extract(path: Path):
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages.append({"page": i + 1, "text": text})
    full = "\n\n".join(p["text"] for p in pages)
    headings = []
    for line in full.splitlines():
        s = line.strip()
        if re.match(r"^第[一二三四五六七八九十]+章", s) or re.match(r"^\d+(\.\d+)*\s+", s) or s in {"摘 要", "目录", "参考文献", "致谢"}:
            headings.append(s)
    return {"page_count": len(pages), "headings": headings[:200], "full_text": full}


def project_inventory(root: Path):
    files = [p for p in root.rglob("*") if p.is_file() and ".git" not in p.parts]
    suffixes = Counter(p.suffix.lower() or "<none>" for p in files)
    java_files = [p for p in files if p.suffix == ".java"]
    packages = Counter()
    classes = []
    for p in java_files:
        text = p.read_text(encoding="utf-8", errors="ignore")
        m = re.search(r"package\s+([\w.]+);", text)
        if m:
            packages[m.group(1)] += 1
        cm = re.search(r"(class|interface|record|enum)\s+(\w+)", text)
        if cm:
            classes.append({
                "file": str(p.relative_to(root)),
                "kind": cm.group(1),
                "name": cm.group(2),
                "lines": text.count("\n") + 1,
            })
    return {
        "file_count": len(files),
        "suffixes": suffixes.most_common(),
        "java_count": len(java_files),
        "packages": packages.most_common(),
        "classes": classes,
    }


def main():
    template = docx_summary(TEMPLATE)
    pdf = pdf_extract(PDF)
    inv = project_inventory(PROJECT)
    (OUT / "template_summary.json").write_text(json.dumps(template, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT / "reference_report_text.txt").write_text(pdf["full_text"], encoding="utf-8")
    (OUT / "reference_report_outline.json").write_text(json.dumps({k: v for k, v in pdf.items() if k != "full_text"}, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT / "project_inventory.json").write_text(json.dumps(inv, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "template_paragraphs": len(template["paragraphs"]),
        "template_tables": len(template["tables"]),
        "reference_pages": pdf["page_count"],
        "outline_preview": pdf["headings"][:30],
        "project_files": inv["file_count"],
        "java_files": inv["java_count"],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
